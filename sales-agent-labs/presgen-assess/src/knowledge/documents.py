"""Document processing pipeline for PresGen-Assess."""

import hashlib
import logging
import mimetypes
from pathlib import Path
from typing import Dict, List, Optional, Tuple
from uuid import uuid4

import PyPDF2
from docx import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Handles document processing with content classification and semantic chunking."""

    def __init__(self):
        """Initialize document processor with chunking strategy."""
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
            separators=["\n\n", "\n", ". ", " ", ""]
        )

        # Supported file types
        self.supported_types = {
            'application/pdf': self._extract_pdf_text,
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': self._extract_docx_text,
            'text/plain': self._extract_txt_text,
            'text/markdown': self._extract_txt_text,  # Treat markdown as plain text
            'text/x-markdown': self._extract_txt_text  # Alternative markdown MIME type
        }

    async def process_certification_document(
        self,
        file_path: Path,
        certification_id: str,
        content_classification: str = "exam_guide",
        original_filename: Optional[str] = None
    ) -> Dict:
        """Process uploaded certification materials into chunks with metadata."""
        try:
            # Validate file
            if not file_path.exists():
                raise FileNotFoundError(f"File not found: {file_path}")

            # Detect file type
            mime_type, _ = mimetypes.guess_type(str(file_path))
            if mime_type not in self.supported_types:
                raise ValueError(f"Unsupported file type: {mime_type}")

            # Calculate file hash for deduplication
            file_hash = self._calculate_file_hash(file_path)

            # Extract text content
            text_content = await self._extract_text_content(file_path, mime_type)

            if not text_content.strip():
                raise ValueError("No text content extracted from document")

            # Generate chunks with semantic splitting
            chunks = self._create_semantic_chunks(text_content)

            # Generate metadata for each chunk
            metadata_list = self._generate_chunk_metadata(
                chunks=chunks,
                certification_id=certification_id,
                content_classification=content_classification,
                original_filename=original_filename or file_path.name,
                file_hash=file_hash,
                mime_type=mime_type
            )

            logger.info(
                f"✅ Processed document: {file_path.name} "
                f"({len(chunks)} chunks, {len(text_content)} chars)"
            )

            return {
                "success": True,
                "chunks": chunks,
                "metadata": metadata_list,
                "file_info": {
                    "original_filename": original_filename or file_path.name,
                    "file_size": file_path.stat().st_size,
                    "file_hash": file_hash,
                    "mime_type": mime_type,
                    "content_classification": content_classification,
                    "chunk_count": len(chunks),
                    "character_count": len(text_content)
                }
            }

        except Exception as e:
            logger.error(f"❌ Failed to process document {file_path}: {e}")
            return {
                "success": False,
                "error": str(e),
                "file_path": str(file_path)
            }

    async def _extract_text_content(self, file_path: Path, mime_type: str) -> str:
        """Extract text content based on file type."""
        extractor = self.supported_types.get(mime_type)
        if not extractor:
            raise ValueError(f"No extractor available for mime type: {mime_type}")

        try:
            return extractor(file_path)
        except Exception as e:
            logger.error(f"❌ Text extraction failed for {file_path}: {e}")
            raise

    def _extract_pdf_text(self, file_path: Path) -> str:
        """Extract text from PDF files."""
        text_content = []

        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)

                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            # Add page marker for better chunking
                            text_content.append(f"\n--- Page {page_num + 1} ---\n{page_text}")
                    except Exception as e:
                        logger.warning(f"⚠️ Failed to extract text from page {page_num + 1}: {e}")
                        continue

                return "\n".join(text_content)

        except Exception as e:
            logger.error(f"❌ PDF extraction failed: {e}")
            raise

    def _extract_docx_text(self, file_path: Path) -> str:
        """Extract text from DOCX files."""
        text_content = []

        try:
            doc = Document(file_path)

            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)

            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(" | ".join(row_text))

            return "\n".join(text_content)

        except Exception as e:
            logger.error(f"❌ DOCX extraction failed: {e}")
            raise

    def _extract_txt_text(self, file_path: Path) -> str:
        """Extract text from plain text files."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    return file.read()
            except Exception as e:
                logger.error(f"❌ Text file extraction failed: {e}")
                raise

    def _create_semantic_chunks(self, text: str) -> List[str]:
        """Create semantic chunks with overlap for better context preservation."""
        try:
            # Clean and normalize text
            cleaned_text = self._clean_text(text)

            # Split into chunks
            chunks = self.text_splitter.split_text(cleaned_text)

            # Filter out very short or empty chunks
            meaningful_chunks = [
                chunk.strip() for chunk in chunks
                if len(chunk.strip()) > 50  # Minimum 50 characters
            ]

            return meaningful_chunks

        except Exception as e:
            logger.error(f"❌ Chunking failed: {e}")
            raise

    def _clean_text(self, text: str) -> str:
        """Clean and normalize text content."""
        # Remove excessive whitespace
        text = " ".join(text.split())

        # Remove page markers but preserve structure
        text = text.replace("--- Page", "\n\n--- Page")

        # Normalize line breaks
        text = text.replace("\r\n", "\n").replace("\r", "\n")

        return text

    def _generate_chunk_metadata(
        self,
        chunks: List[str],
        certification_id: str,
        content_classification: str,
        original_filename: str,
        file_hash: str,
        mime_type: str
    ) -> List[Dict]:
        """Generate comprehensive metadata for each chunk."""
        metadata_list = []

        for i, chunk in enumerate(chunks):
            # Extract potential page reference
            page_num = self._extract_page_number(chunk)

            # Determine domain/subdomain from content (simplified approach)
            domain_info = self._classify_content_domain(chunk)

            metadata = {
                "certification_id": certification_id,
                "content_classification": content_classification,
                "document_name": original_filename,
                "chunk_index": i,
                "chunk_id": str(uuid4()),
                "file_hash": file_hash,
                "mime_type": mime_type,
                "source_page": page_num,
                "chunk_size": len(chunk),
                "domain": domain_info.get("domain", "general"),
                "subdomain": domain_info.get("subdomain"),
                "content_type": self._classify_content_type(chunk),
                "created_at": self._get_current_timestamp()
            }

            metadata_list.append(metadata)

        return metadata_list

    def _extract_page_number(self, chunk: str) -> Optional[int]:
        """Extract page number from chunk content."""
        import re
        page_match = re.search(r"--- Page (\d+) ---", chunk)
        if page_match:
            return int(page_match.group(1))
        return None

    def _classify_content_domain(self, chunk: str) -> Dict[str, Optional[str]]:
        """Basic content domain classification (can be enhanced with ML)."""
        chunk_lower = chunk.lower()

        # Common certification domains
        domain_keywords = {
            "security": ["security", "encryption", "authentication", "authorization", "compliance"],
            "networking": ["network", "vpc", "subnet", "routing", "load balancer"],
            "compute": ["instance", "virtual machine", "container", "serverless", "lambda"],
            "storage": ["storage", "database", "backup", "archive", "data"],
            "monitoring": ["monitoring", "logging", "metrics", "alerting", "cloudwatch"],
            "management": ["management", "governance", "cost", "billing", "organization"]
        }

        for domain, keywords in domain_keywords.items():
            if any(keyword in chunk_lower for keyword in keywords):
                return {"domain": domain, "subdomain": None}

        return {"domain": "general", "subdomain": None}

    def _classify_content_type(self, chunk: str) -> str:
        """Classify the type of content in the chunk."""
        chunk_lower = chunk.lower()

        if any(word in chunk_lower for word in ["definition", "concept", "introduction"]):
            return "concept"
        elif any(word in chunk_lower for word in ["step", "procedure", "process", "how to"]):
            return "procedure"
        elif any(word in chunk_lower for word in ["example", "case study", "scenario"]):
            return "example"
        elif any(word in chunk_lower for word in ["best practice", "recommendation", "guideline"]):
            return "best_practice"
        else:
            return "general"

    def _calculate_file_hash(self, file_path: Path) -> str:
        """Calculate SHA-256 hash of file for deduplication."""
        hash_sha256 = hashlib.sha256()
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hash_sha256.update(chunk)
        return hash_sha256.hexdigest()

    def _get_current_timestamp(self) -> str:
        """Get current timestamp in ISO format."""
        from datetime import datetime
        return datetime.now().isoformat()

    def validate_file(self, file_path: Path, max_size_mb: int = 100) -> Tuple[bool, Optional[str]]:
        """Validate file before processing."""
        try:
            # Check if file exists
            if not file_path.exists():
                return False, "File does not exist"

            # Check file size
            file_size_mb = file_path.stat().st_size / (1024 * 1024)
            if file_size_mb > max_size_mb:
                return False, f"File size ({file_size_mb:.1f}MB) exceeds limit ({max_size_mb}MB)"

            # Check file type
            mime_type, _ = mimetypes.guess_type(str(file_path))
            if mime_type not in self.supported_types:
                return False, f"Unsupported file type: {mime_type}"

            return True, None

        except Exception as e:
            return False, f"Validation error: {str(e)}"

    def get_processing_stats(self, processing_result: Dict) -> Dict:
        """Get processing statistics from result."""
        if not processing_result.get("success"):
            return {"status": "failed", "error": processing_result.get("error")}

        file_info = processing_result.get("file_info", {})
        return {
            "status": "success",
            "file_size_mb": round(file_info.get("file_size", 0) / (1024 * 1024), 2),
            "chunk_count": file_info.get("chunk_count", 0),
            "character_count": file_info.get("character_count", 0),
            "content_classification": file_info.get("content_classification"),
            "mime_type": file_info.get("mime_type"),
            "average_chunk_size": round(
                file_info.get("character_count", 0) / max(file_info.get("chunk_count", 1), 1)
            )
        }